#!/usr/bin/env python
"""Extract evidence from experiment report DOCX files.

This script is evidence-only. It never decides deductions, final scores,
comments, plagiarism, formatting penalties, or experiment completeness.

Design goals:
- Work across courses with folder names like "experiment1", numeric Chinese
  experiment folders, and ordinal experiment folders.
- Skip manuals, gradebooks, final/regraded outputs, and temporary Word locks by
  default.
- Keep all filtering configurable so a project does not need an extra cleanup
  script.
"""

from __future__ import annotations

import argparse
import csv
import re
import zipfile
from pathlib import Path

from docx import Document


C = {
    "exp": "\u5b9e\u9a8c",
    "sid": "\u5b66\u53f7",
    "name": "\u59d3\u540d",
    "file": "\u6587\u4ef6",
    "late": "\u662f\u5426\u8865\u4ea4",
    "image_count": "\u56fe\u7247\u6570",
    "content_len": "\u5b9e\u9a8c\u5185\u5bb9\u957f\u5ea6",
    "process_len": "\u5b9e\u9a8c\u8fc7\u7a0b\u957f\u5ea6",
    "summary_len": "\u603b\u7ed3\u957f\u5ea6",
    "content_preview": "\u5b9e\u9a8c\u5185\u5bb9\u6458\u5f55",
    "process_preview": "\u5b9e\u9a8c\u8fc7\u7a0b\u6458\u5f55",
    "summary_preview": "\u603b\u7ed3\u6458\u5f55",
    "skip_reason": "\u8df3\u8fc7\u539f\u56e0",
    "read_error": "\u8bfb\u53d6\u9519\u8bef",
    "script_note": "\u811a\u672c\u7ed3\u8bba",
}

YES = "\u662f"
NO = "\u5426"
NOTE = "\u8bc1\u636e\u63d0\u53d6\uff0c\u4ec5\u4f9b\u4eba\u5de5\u590d\u6838"

DEFAULT_SKIP_KEYWORDS = [
    "\u5b9e\u9a8c\u624b\u518c",
    "\u624b\u518c",
    "\u6307\u5bfc\u4e66",
    "\u4efb\u52a1\u4e66",
    "\u8bc4\u5206\u7ec6\u5219",
    "\u6210\u7ee9\u8868",
    "\u6210\u7ee9\u6c47\u603b",
    "\u767b\u5206\u8868",
]
DEFAULT_REVIEWED_KEYWORDS = [
    "\u6700\u7ec8\u6279\u9605\u7248",
    "\u91cd\u65b0\u6279\u9605\u7248",
    "\u91cd\u65b0\u6279\u6539\u7248",
]
LATE_KEYWORDS = ["\u8865\u4ea4", "\u8fdf\u4ea4", "\u903e\u671f"]

CONTENT_STARTS = ["\u5b9e\u9a8c\u5185\u5bb9", "\u5b9e\u9a8c\u8981\u6c42", "\u5b9e\u9a8c\u4efb\u52a1"]
PROCESS_STARTS = [
    "\u5b9e\u9a8c\u6b65\u9aa4\u3001\u5b9e\u9a8c\u7ed3\u679c\u53ca\u5206\u6790",
    "\u5b9e\u9a8c\u6b65\u9aa4",
    "\u5b9e\u9a8c\u8fc7\u7a0b",
    "\u64cd\u4f5c\u6b65\u9aa4",
    "\u5b9e\u9a8c\u7ed3\u679c\u53ca\u5206\u6790",
]
SUMMARY_STARTS = ["\u5b9e\u9a8c\u603b\u7ed3", "\u603b\u7ed3", "\u5fc3\u5f97\u4f53\u4f1a"]
COMMON_ENDS = [
    "\u5b9e\u9a8c\u8bbe\u5907",
    "\u5b9e\u9a8c\u73af\u5883",
    "\u5b9e\u9a8c\u6b65\u9aa4",
    "\u5b9e\u9a8c\u8fc7\u7a0b",
    "\u5b9e\u9a8c\u7ed3\u679c",
    "\u5b9e\u9a8c\u603b\u7ed3",
    "\u6559\u5e08\u8bc4\u8bed",
    "\u6210\u7ee9\u8bc4\u5b9a",
]


CN_DIGITS = {
    "\u96f6": 0,
    "\u4e00": 1,
    "\u4e8c": 2,
    "\u4e24": 2,
    "\u4e09": 3,
    "\u56db": 4,
    "\u4e94": 5,
    "\u516d": 6,
    "\u4e03": 7,
    "\u516b": 8,
    "\u4e5d": 9,
}


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def fullwidth_to_ascii(text: str) -> str:
    return text.translate(str.maketrans("０１２３４５６７８９", "0123456789"))


def chinese_num_to_int(text: str) -> int | None:
    text = text.strip()
    if not text:
        return None
    if text.isdigit():
        return int(text)
    ten = "\u5341"
    if ten not in text:
        if len(text) == 1:
            return CN_DIGITS.get(text)
        total = 0
        for ch in text:
            if ch not in CN_DIGITS:
                return None
            total = total * 10 + CN_DIGITS[ch]
        return total
    left, _, right = text.partition(ten)
    tens = 1 if not left else CN_DIGITS.get(left)
    ones = 0 if not right else CN_DIGITS.get(right)
    if tens is None or ones is None:
        return None
    return tens * 10 + ones


def detect_experiment(path: Path) -> str:
    candidates = []
    joined_parts = list(path.parts)
    patterns = [
        re.compile(r"(?:\u7b2c)?\s*([0-9\uff10-\uff19]+|[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u4e24]+)\s*(?:\u6b21)?\s*\u5b9e\u9a8c"),
        re.compile(r"\u5b9e\u9a8c\s*([0-9\uff10-\uff19]+|[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u4e24]+)"),
        re.compile(r"exp(?:eriment)?\s*([0-9]+)", re.I),
    ]
    for part in reversed(joined_parts):
        part_norm = fullwidth_to_ascii(part)
        for pattern in patterns:
            match = pattern.search(part_norm)
            if match:
                value = chinese_num_to_int(fullwidth_to_ascii(match.group(1)))
                if value is not None:
                    candidates.append(value)
    return str(candidates[0]) if candidates else ""


def text_from_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                if cell_text.strip():
                    parts.append(cell_text)
    return "\n".join(parts)


def section(text: str, starts: list[str], ends: list[str]) -> str:
    positions = [(text.find(s), s) for s in starts if text.find(s) >= 0]
    if not positions:
        return ""
    pos, label = min(positions, key=lambda x: x[0])
    end = len(text)
    for e in ends:
        p = text.find(e, pos + len(label))
        if p >= 0:
            end = min(end, p)
    return normalize(text[pos:end])


def image_count(path: Path) -> int:
    try:
        with zipfile.ZipFile(path) as zf:
            return sum(1 for n in zf.namelist() if n.startswith("word/media/"))
    except Exception:
        return 0


def parse_identity(path: Path) -> tuple[str, str]:
    stem = re.sub(r"[（(].*?[）)]", "", path.stem)
    stem = fullwidth_to_ascii(stem)
    match = re.match(r"(?P<sid>\d{8,})\s*[-_－]?\s*(?P<name>.*)", stem)
    if not match:
        return "", stem.strip()
    return match.group("sid"), match.group("name").strip("-_－ ")


def should_skip(path: Path, args: argparse.Namespace) -> str:
    name = path.name
    full = str(path)
    if name.startswith("~$"):
        return "temporary Word lock file"
    if args.skip_manuals and any(k in full for k in args.skip_keyword):
        return "manual/task/gradebook keyword"
    if not args.include_reviewed and any(k in name for k in DEFAULT_REVIEWED_KEYWORDS):
        return "reviewed output"
    if args.exclude_name_regex and re.search(args.exclude_name_regex, full):
        return "excluded by regex"
    if args.include_name_regex and not re.search(args.include_name_regex, full):
        return "not matched by include regex"
    return ""


def report_rows(root: Path, args: argparse.Namespace) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    rows = []
    skipped = []
    for path in sorted(root.rglob("*.docx")):
        reason = should_skip(path, args)
        if reason:
            skipped.append({C["file"]: str(path.resolve()), C["skip_reason"]: reason})
            continue
        sid, name = parse_identity(path)
        exp = detect_experiment(path)
        late = YES if any(x in path.name for x in LATE_KEYWORDS) else NO
        try:
            text = text_from_docx(path)
            content = section(text, CONTENT_STARTS, COMMON_ENDS)
            process = section(text, PROCESS_STARTS, ["\u5b9e\u9a8c\u603b\u7ed3", "\u6559\u5e08\u8bc4\u8bed", "\u6210\u7ee9\u8bc4\u5b9a"])
            summary = section(text, SUMMARY_STARTS, ["\u6559\u5e08\u8bc4\u8bed", "\u6210\u7ee9\u8bc4\u5b9a"])
            error = ""
        except Exception as exc:
            content = process = summary = ""
            error = str(exc)
        rows.append(
            {
                C["exp"]: exp,
                C["sid"]: sid,
                C["name"]: name,
                C["file"]: str(path.resolve()),
                C["late"]: late,
                C["image_count"]: str(image_count(path)),
                C["content_len"]: str(len(content)),
                C["process_len"]: str(len(process)),
                C["summary_len"]: str(len(summary)),
                C["content_preview"]: content[: args.content_preview],
                C["process_preview"]: process[: args.process_preview],
                C["summary_preview"]: summary[: args.summary_preview],
                C["read_error"]: error,
                C["script_note"]: NOTE,
            }
        )
    return rows, skipped


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract evidence from experiment report DOCX files.")
    parser.add_argument("root", help="Course/report root directory")
    parser.add_argument("--out", default="report_evidence.csv", help="Output evidence CSV path")
    parser.add_argument("--skipped-out", default="", help="Optional skipped-files CSV path")
    parser.add_argument("--include-reviewed", action="store_true", help="Include final/regraded reviewed output files")
    parser.add_argument("--skip-manuals", action=argparse.BooleanOptionalAction, default=True, help="Skip manuals/task books/gradebooks by keyword")
    parser.add_argument("--skip-keyword", action="append", default=list(DEFAULT_SKIP_KEYWORDS), help="Extra keyword to skip; may be repeated")
    parser.add_argument("--include-name-regex", default="", help="Only include paths matching this regex")
    parser.add_argument("--exclude-name-regex", default="", help="Exclude paths matching this regex")
    parser.add_argument("--content-preview", type=int, default=500)
    parser.add_argument("--process-preview", type=int, default=1000)
    parser.add_argument("--summary-preview", type=int, default=800)
    args = parser.parse_args()

    rows, skipped = report_rows(Path(args.root), args)
    fields = [
        C["exp"],
        C["sid"],
        C["name"],
        C["file"],
        C["late"],
        C["image_count"],
        C["content_len"],
        C["process_len"],
        C["summary_len"],
        C["content_preview"],
        C["process_preview"],
        C["summary_preview"],
        C["read_error"],
        C["script_note"],
    ]
    write_csv(Path(args.out), rows, fields)
    if args.skipped_out:
        write_csv(Path(args.skipped_out), skipped, [C["file"], C["skip_reason"]])
    print(f"wrote {len(rows)} evidence rows to {args.out}")
    print(f"skipped {len(skipped)} files")


if __name__ == "__main__":
    main()
