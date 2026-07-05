#!/usr/bin/env python
"""Generate plagiarism-similarity candidates for manual review.

The output is only a review queue. It must not be treated as a plagiarism
decision, final deduction, or final score.
"""

from __future__ import annotations

import argparse
import csv
import itertools
import re
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document


REMOVE_RE = re.compile(
    "|".join(
        [
            "实验报告",
            "实验名称",
            "实验目的",
            "实验内容",
            "实验步骤",
            "实验结果及分析",
            "实验总结",
            "心得体会",
            "过程总结",
            "教师评语",
            "成绩评定",
            "教师签字",
        ]
    )
)
COMMENT_RE = re.compile(r"教师评语.*?(成绩评定|教师签字|$)", re.S)


def normalize(text: str) -> str:
    text = COMMENT_RE.sub(" ", text)
    text = REMOVE_RE.sub(" ", text)
    return re.sub(r"\s+", "", text)


def doc_text(path: Path) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                if txt:
                    parts.append(txt)
    return normalize("\n".join(parts))


def shingles(text: str, n: int) -> set[str]:
    if len(text) < n:
        return {text} if text else set()
    return {text[i : i + n] for i in range(len(text) - n + 1)}


def longest_common(a: str, b: str) -> tuple[int, str]:
    match = max(SequenceMatcher(None, a, b, autojunk=False).get_matching_blocks(), key=lambda m: m.size)
    return match.size, a[match.a : match.a + min(match.size, 260)]


def load_rows(evidence_csv: Path) -> list[dict[str, str]]:
    with evidence_csv.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    for row in rows:
        row["_text"] = doc_text(Path(row["文件"]))
        row["_shingles"] = shingles(row["_text"], 9)
    return rows


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a manual-review plagiarism candidate CSV.")
    parser.add_argument("evidence_csv", help="CSV produced by extract_report_evidence.py")
    parser.add_argument("--out", default="plagiarism_candidates.csv")
    parser.add_argument("--jaccard", type=float, default=0.50)
    parser.add_argument("--containment", type=float, default=0.75)
    args = parser.parse_args()

    rows = load_rows(Path(args.evidence_csv))
    candidates = []
    for exp in sorted({r.get("实验", "") for r in rows}):
        group = [r for r in rows if r.get("实验", "") == exp]
        for a, b in itertools.combinations(group, 2):
            sa, sb = a["_shingles"], b["_shingles"]
            if not sa or not sb:
                continue
            inter = len(sa & sb)
            jac = inter / len(sa | sb)
            cont = max(inter / len(sa), inter / len(sb))
            if jac >= args.jaccard or cont >= args.containment:
                length, fragment = longest_common(a["_text"], b["_text"])
                candidates.append(
                    {
                        "实验": exp,
                        "学号A": a.get("学号", ""),
                        "姓名A": a.get("姓名", ""),
                        "文件A": a.get("文件", ""),
                        "学号B": b.get("学号", ""),
                        "姓名B": b.get("姓名", ""),
                        "文件B": b.get("文件", ""),
                        "文本Jaccard": f"{jac:.3f}",
                        "文本包含度": f"{cont:.3f}",
                        "最长连续重合字数": str(length),
                        "最长重合片段": fragment,
                        "脚本结论": "疑似雷同，待人工复核；不得直接判定抄袭",
                    }
                )
    candidates.sort(key=lambda r: (float(r["文本Jaccard"]), float(r["文本包含度"])), reverse=True)
    fields = [
        "实验",
        "学号A",
        "姓名A",
        "文件A",
        "学号B",
        "姓名B",
        "文件B",
        "文本Jaccard",
        "文本包含度",
        "最长连续重合字数",
        "最长重合片段",
        "脚本结论",
    ]
    with open(args.out, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(candidates)
    print(f"wrote {len(candidates)} candidates to {args.out}")


if __name__ == "__main__":
    main()
