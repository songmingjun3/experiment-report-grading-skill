#!/usr/bin/env python
"""Write manually reviewed grades/comments to Word reports and an Excel gradebook.

Input must be a manual decision CSV. This script refuses to run unless the CSV
contains an explicit manual-basis column. It does not calculate scores.

Safety rule: only replace the grading-comment paragraph and score paragraph in a
recognized grading cell. If the Word structure is abnormal, record the row in an
exception CSV and skip Word writing for that report. Never append comments to the
end of the document or clear the whole grading cell.
"""

from __future__ import annotations

import argparse
import csv
import re
import shutil
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Pt

EXP = "\u5b9e\u9a8c"
SID = "\u5b66\u53f7"
NAME = "\u59d3\u540d"
SRC = "\u6e90\u6587\u4ef6"
SCORE = "\u6700\u7ec8\u6210\u7ee9"
COMMENT = "\u6700\u7ec8\u8bc4\u8bed"
BASIS = "\u4eba\u5de5\u5224\u5b9a\u4f9d\u636e"
FINAL_REPORT = "\u6700\u7ec8\u62a5\u544a"
WORD_STATUS = "Word\u5199\u5165\u72b6\u6001"
WORD_ERROR = "Word\u5199\u5165\u5f02\u5e38"
MARK = "\u6807\u8bb0"
IS_MAKEUP = "\u662f\u5426\u8865\u4ea4"
GRADE_LABEL = "\u6210\u7ee9\u8bc4\u5b9a"
POINT_LABEL = "\u5f97\u5206"
TEACHER_COMMENT = "\u6307\u5bfc\u6559\u5e08\u8bc4\u8bed"
CHINESE_NUMS = {
    1: "\u4e00",
    2: "\u4e8c",
    3: "\u4e09",
    4: "\u56db",
    5: "\u4e94",
    6: "\u516d",
    7: "\u4e03",
    8: "\u516b",
    9: "\u4e5d",
    10: "\u5341",
    11: "\u5341\u4e00",
    12: "\u5341\u4e8c",
}
COPY_FLAG = "\u6284\u88ad"
MAKEUP_FLAG = "\u8865\u4ea4"
YES = "\u662f"
SUCCESS = "\u6210\u529f"
SKIPPED = "\u8df3\u8fc7"
EXCEPTION_CSV = "\u5199\u5165\u5f02\u5e38\u8bb0\u5f55.csv"

REQUIRED_COLUMNS = [EXP, SID, NAME, SRC, SCORE, COMMENT, BASIS]


def clean_stem(path: Path) -> str:
    return re.sub(r"[\uFF08(].*?\u6279\u9605.*?[\uFF09)]", "", path.stem).strip()


def clear_existing_score(paragraph) -> None:
    seen = False
    for run in paragraph.runs:
        text = run.text or ""
        if GRADE_LABEL in text or POINT_LABEL in text:
            seen = True
            continue
        if seen:
            run.text = re.sub(r"\d{1,3}", "", text).replace("\u5206", "")


def locate_grading_paragraphs(doc: Document):
    candidates = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in {id(c._tc) for c in candidates}:
                    continue
                if TEACHER_COMMENT in cell.text and GRADE_LABEL in cell.text:
                    candidates.append(cell)
    if not candidates:
        if not doc.tables:
            raise ValueError("missing first table")
        raise ValueError("cannot locate grading cell with teacher comment and score labels")
    if len(candidates) > 1:
        raise ValueError("multiple possible grading cells; skip to avoid wrong write")
    cell = candidates[0]
    paragraphs = cell.paragraphs
    label_indexes = [i for i, p in enumerate(paragraphs) if TEACHER_COMMENT in p.text]
    score_indexes = [i for i, p in enumerate(paragraphs) if GRADE_LABEL in p.text]
    if not label_indexes:
        raise ValueError("grading cell does not contain teacher comment label")
    if not score_indexes:
        raise ValueError("grading cell lacks score paragraph")
    label_i = label_indexes[0]
    score_i = score_indexes[0]
    if score_i <= label_i:
        raise ValueError("score paragraph appears before teacher comment label")
    comment_paragraphs = paragraphs[label_i + 1:score_i]
    if not comment_paragraphs:
        raise ValueError("grading cell lacks comment body paragraph")
    return comment_paragraphs, paragraphs[score_i]

def write_docx(src: Path, dst: Path, score: str, comment: str) -> None:
    doc = Document(src)
    comment_paragraphs, score_p = locate_grading_paragraphs(doc)
    first = comment_paragraphs[0]
    for paragraph in comment_paragraphs:
        paragraph.clear()
    run = first.add_run(comment)
    run.font.size = Pt(12)
    clear_existing_score(score_p)
    replaced = False
    for run in score_p.runs:
        if run.font.underline and run.text.strip() not in {GRADE_LABEL, POINT_LABEL}:
            run.text = f"  {score}  "
            replaced = True
            break
    if not replaced:
        score_p.add_run(f"  {score}  ")
    doc.save(dst)

def validate_rows(rows: list[dict[str, str]]) -> None:
    if not rows:
        raise SystemExit("manual decision CSV is empty")
    missing = [c for c in REQUIRED_COLUMNS if c not in rows[0]]
    if missing:
        raise SystemExit(f"manual decision CSV missing required columns: {', '.join(missing)}")
    for i, row in enumerate(rows, 2):
        if not row.get(BASIS, "").strip():
            raise SystemExit(f"row {i} missing manual basis; refuse to write final results")


def normalize_id(value: object) -> str:
    text = str(value or "").strip()
    if text.endswith(".0"):
        text = text[:-2]
    return re.sub(r"\s+", "", text)


def normalize_header(value: object) -> str:
    return re.sub(r"\s+", "", str(value or "").strip())


def experiment_header_candidates(exp: int) -> list[str]:
    cn = CHINESE_NUMS.get(exp)
    candidates = [f"{EXP}{exp}", f"{EXP}{exp:02d}", f"\u7b2c{exp}\u6b21\u5b9e\u9a8c"]
    if cn:
        candidates.extend([f"{EXP}{cn}", f"\u7b2c{cn}\u6b21\u5b9e\u9a8c"])
    return [normalize_header(x) for x in candidates]


def resolve_output_path(out_dir: Path, value: str) -> Path:
    path = Path(value)
    if path.is_absolute():
        return path
    if path.parent == Path("."):
        return out_dir / path
    return path


def update_gradebook(template: Path, output: Path, rows: list[dict[str, str]]) -> int:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    wb = openpyxl.load_workbook(output)
    ws = wb.active
    headers = {normalize_header(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1)}
    sid_col = headers.get(normalize_header(SID))
    name_col = headers.get(normalize_header(NAME))
    if not sid_col:
        raise ValueError("gradebook missing student id column")
    by_sid = {(str(r[EXP]).strip(), normalize_id(r[SID])): r for r in rows}
    by_name: dict[tuple[str, str], dict[str, str] | None] = {}
    for row in rows:
        key = (str(row[EXP]).strip(), normalize_header(row.get(NAME, "")))
        by_name[key] = row if key not in by_name else None
    exp_cols: dict[int, int] = {}
    for exp in range(1, 50):
        for candidate in experiment_header_candidates(exp):
            if candidate in headers:
                exp_cols[exp] = headers[candidate]
                break
    red_font = Font(color="FFFF0000")
    red_fill = PatternFill(fill_type="solid", fgColor="FFFFE5E5")
    blue_font = Font(color="FF0000FF")
    blue_fill = PatternFill(fill_type="solid", fgColor="FFE5F0FF")
    updated = 0
    for rr in range(2, ws.max_row + 1):
        sid = normalize_id(ws.cell(rr, sid_col).value)
        name = normalize_header(ws.cell(rr, name_col).value) if name_col else ""
        for exp, col in exp_cols.items():
            row = by_sid.get((str(exp), sid))
            if not row and name:
                row = by_name.get((str(exp), name))
            if not row:
                continue
            cell = ws.cell(rr, col)
            cell.value = int(float(row[SCORE]))
            flag = row.get(MARK, "")
            if COPY_FLAG in flag:
                cell.font = blue_font
                cell.fill = blue_fill
            elif row.get(IS_MAKEUP) == YES or MAKEUP_FLAG in flag:
                cell.font = red_font
                cell.fill = red_fill
            updated += 1
    wb.save(output)
    wb.close()
    return updated

def main() -> None:
    parser = argparse.ArgumentParser(description="Write manual experiment report grading results.")
    parser.add_argument("manual_csv", help="Manual decision CSV with required columns")
    parser.add_argument("--out-dir", required=True, help="Output directory for reviewed reports and summaries")
    parser.add_argument("--gradebook", help="Optional gradebook template .xlsx")
    parser.add_argument("--gradebook-out", default="final_gradebook.xlsx")
    parser.add_argument("--exceptions-out", default=EXCEPTION_CSV, help="CSV for skipped Word writes")
    args = parser.parse_args()

    with open(args.manual_csv, "r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    validate_rows(rows)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    ok = 0
    exceptions = []
    for row in rows:
        src = Path(row[SRC])
        exp_dir = out_dir / f"{EXP}{row[EXP]}\u6700\u7ec8\u6279\u9605\u7248"
        exp_dir.mkdir(exist_ok=True)
        dst = exp_dir / f"{clean_stem(src)}\uFF08\u6700\u7ec8\u6279\u9605\u7248\uFF09.docx"
        try:
            write_docx(src, dst, row[SCORE], row[COMMENT])
        except Exception as exc:
            row[FINAL_REPORT] = ""
            row[WORD_STATUS] = SKIPPED
            row[WORD_ERROR] = str(exc)
            exceptions.append(
                {
                    EXP: row.get(EXP, ""),
                    SID: row.get(SID, ""),
                    NAME: row.get(NAME, ""),
                    SRC: row.get(SRC, ""),
                    "\u5f02\u5e38\u539f\u56e0": str(exc),
                    "\u5904\u7406": "\u5df2\u8df3\u8fc7 Word \u5199\u5165\uff1b\u672a\u8ffd\u52a0\u5230\u6587\u672b\uff1b\u672a\u6539\u5199\u8868\u683c",
                }
            )
            continue
        row[FINAL_REPORT] = str(dst.resolve())
        row[WORD_STATUS] = SUCCESS
        row[WORD_ERROR] = ""
        ok += 1
    summary_csv = out_dir / "\u6700\u7ec8\u4eba\u5de5\u8bc4\u5206\u660e\u7ec6.csv"
    fields = list(rows[0].keys())
    for extra in [FINAL_REPORT, WORD_STATUS, WORD_ERROR]:
        if extra not in fields:
            fields.append(extra)
    with summary_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)
    exceptions_csv = resolve_output_path(out_dir, args.exceptions_out)
    with exceptions_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[EXP, SID, NAME, SRC, "\u5f02\u5e38\u539f\u56e0", "\u5904\u7406"],
        )
        writer.writeheader()
        writer.writerows(exceptions)
    updated = 0
    if args.gradebook:
        updated = update_gradebook(Path(args.gradebook), resolve_output_path(out_dir, args.gradebook_out), rows)
    print(f"wrote {ok} reviewed reports")
    print(f"skipped {len(exceptions)} reports with Word structure exceptions")
    print(f"wrote {summary_csv}")
    print(f"wrote {exceptions_csv}")
    if args.gradebook:
        print(f"updated {updated} gradebook cells")


if __name__ == "__main__":
    main()



